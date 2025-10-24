import os
from dotenv import load_dotenv
load_dotenv()
import gradio as gr
import google.generativeai as genai
import sys
import json
from PIL import Image
import PyPDF2
from docx import Document
import mimetypes

sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from rag_utils import prepare_rag_store, retrieve

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise RuntimeError("Set the GOOGLE_API_KEY environment variable.")

genai.configure(api_key=GOOGLE_API_KEY)
gemini_model = genai.GenerativeModel("gemini-2.5-flash")


json_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "visa_rules.json")
try:
    with open(json_path, 'r', encoding='utf-8') as f:
        rules = json.load(f)
    print(f"Loaded {len(rules)} rules from {json_path}")
except FileNotFoundError:
    print(f"Warning: {json_path} not found. Using empty rules.")
    rules = []


try:
    rules_rag, embeddings, index, model = prepare_rag_store()
    if len(rules_rag) > len(rules):
        rules = rules_rag
        print(f"Using {len(rules)} rules from RAG store")
except Exception as e:
    print(f"Warning: Could not prepare RAG store: {e}")
    embeddings, index, model = None, None, None

def get_country_categories(country):
    if not country:
        return []
    cats = sorted(set(rule["category"] for rule in rules if rule.get("country") and rule["country"].lower() == country.lower()))
    cats = [c.title() for c in cats if c and c.strip()]
    return cats

def get_category_requirements(country, category):
    for rule in rules:
        rule_country = rule.get("country", "")
        rule_category = rule.get("category", "")
        if rule_country and rule_country.lower() == country.lower() and rule_category and rule_category.lower() == category.lower():
            if "requirements" in rule and isinstance(rule["requirements"], list):
                return rule["requirements"]
            req_text = rule.get("text", "")
            if "Required:" in req_text:
                reqs = req_text.split("Required:")[1].split(",")
            elif "Requirements:" in req_text:
                reqs = req_text.split("Requirements:")[1].split(",")
            else:
                reqs = []
            reqs = [r.strip().rstrip('.') for r in reqs if r.strip()]
            return reqs
    return []

def get_rag_context(country, category):
    """Get RAG context for the selected country and category"""
    if not country or not category:
        return "### 📋 RAG Context\n\nPlease select both **country** and **category** to view relevant visa requirements."
    
    
    filtered_docs = [
        rule for rule in rules 
        if rule.get("country", "").lower() == country.lower() 
        and rule.get("category", "").lower() == category.lower()
    ]
    
    
    if not filtered_docs and index is not None and model is not None:
        try:
            search_query = {
                "country": country,
                "category": category.lower(),
                "answers": {}
            }
            retrieved_docs = retrieve(search_query, rules, index, model)
            
            filtered_docs = [
                doc for doc in retrieved_docs 
                if doc.get("country", "").lower() == country.lower() 
                and doc.get("category", "").lower() == category.lower()
            ]
            
            if not filtered_docs:
                filtered_docs = [
                    doc for doc in retrieved_docs 
                    if doc.get("country", "").lower() == country.lower()
                ]
        except Exception as e:
            print(f"RAG retrieval error: {e}")
    
    if not filtered_docs:
        return f"### 📋 RAG Context\n\n**{country} - {category.title()}**\n\nNo specific requirements found.\n\n*Debug: Total rules loaded: {len(rules)}*"
    
    
    context_md = f"### 📋 RAG Context\n\n**{country} - {category.title()}**\n\n"
    for i, doc in enumerate(filtered_docs, 1):
        doc_title = doc.get("title", "Unknown")
        doc_text = doc.get("text", "No details available")
        
        
        if "requirements" in doc and isinstance(doc["requirements"], list):
            context_md += f"#### {doc_title}\n\n"
            context_md += "**Requirements:**\n"
            for req in doc["requirements"]:
                context_md += f"- {req}\n"
            context_md += f"\n**Details:** {doc_text}\n\n"
        else:
            context_md += f"#### {doc_title}\n\n{doc_text}\n\n"
        
        if i < len(filtered_docs):
            context_md += "---\n\n"
    
    return context_md

def build_questions(country, category):
    """Build questions dynamically from requirements or use general visa questions"""
    reqs = get_category_requirements(country, category)
    
    if reqs:
        questions = []
        for i, req in enumerate(reqs):
            req_clean = req.strip()
            questions.append((f"req_{i}", f"Do you have: {req_clean}?"))
        return questions
    else:
        return [
            ("passport_valid", "Do you have a valid passport (valid for at least 6 months)?"),
            ("financial_proof", "Can you provide proof of sufficient funds (bank statements, payslips)?"),
            ("travel_history", "Do you have a good travel history (previous visas to other countries)?"),
            ("employment_status", "Are you currently employed/studying/retired with proof?"),
            ("purpose_clear", "Do you have a clear purpose of visit with supporting documents?"),
            ("ties_home", "Do you have strong ties to your home country (property, family, job)?"),
            ("criminal_record", "Do you have a clean criminal record with no visa rejections?"),
            ("health_insurance", "Do you have travel/health insurance coverage?"),
            ("accommodation", "Do you have confirmed accommodation/invitation letter?"),
            ("return_ticket", "Do you have a return ticket or travel itinerary?")
        ]

def format_prompt(answers, retrieved_docs):
    """Format the prompt with country/category context and retrieved docs"""
    country = answers.get("country", "")
    category = answers.get("category", "")
    
    if retrieved_docs:
        context_parts = []
        for doc in retrieved_docs:
            doc_country = doc.get("country", "Unknown")
            doc_category = doc.get("category", "Unknown")
            doc_text = doc.get("text", "")
            context_parts.append(f"[{doc_country} - {doc_category}]\n{doc_text}")
        context = "\n\n".join(context_parts)
    else:
        context = ""
    
    answer_text = "\n".join([f"- {k}: {v}" for k, v in answers.items() if not k.endswith("_detail")])
    
    if context:
        prompt = (
            f"You are a visa application expert. Use the following official visa requirements as context:\n\n"
            f"{context}\n\n"
            f"IMPORTANT: The applicant is applying for a {category.upper()} visa to {country.upper()}. "
            f"Only use requirements and information specific to {country} {category} visa.\n\n"
            f"Visa Application Information:\n"
            f"{answer_text}\n\n"
            f"Based on this information:\n"
            f"1. Estimate the visa approval probability (as a percentage)\n"
            f"2. List any missing documents or risk factors\n"
            f"3. Provide professional advice to improve approval chances\n"
            f"4. Be specific to {country} {category} visa requirements"
        )
    else:
        prompt = (
            f"You are a visa application expert. Based on your knowledge:\n\n"
            f"The applicant is applying for a {category.upper()} visa to {country.upper()}.\n\n"
            f"Visa Application Information:\n"
            f"{answer_text}\n\n"
            f"Based on this information:\n"
            f"1. Estimate the visa approval probability (as a percentage)\n"
            f"2. List any missing documents or risk factors\n"
            f"3. Provide professional advice to improve approval chances\n"
            f"4. Be specific to {country} {category} visa requirements"
        )
    return prompt

def process_application(selected_country, selected_category, *question_answers):
    """Process the visa application after all questions are answered"""
    questions = build_questions(selected_country, selected_category.lower())
    answers = {}
    answers["country"] = selected_country
    answers["category"] = selected_category.lower()
    
    for i, answer in enumerate(question_answers):
        if i < len(questions) and answer:
            key, question_text = questions[i]
            answers[key] = answer
            answers[f"{key}_detail"] = f"{question_text} Answer: {answer}"
    
    search_query = {
        "country": selected_country,
        "category": selected_category.lower(),
        "answers": answers
    }
    
    retrieved_docs = retrieve(search_query, rules, index, model)
    
    filtered_docs = [
        doc for doc in retrieved_docs 
        if doc.get("country", "").lower() == selected_country.lower() 
        and doc.get("category", "").lower() == selected_category.lower()
    ]
    
    if not filtered_docs:
        filtered_docs = [
            doc for doc in retrieved_docs 
            if doc.get("country", "").lower() == selected_country.lower()
        ]
    
    prompt = format_prompt(answers, filtered_docs)
    
    try:
        response = gemini_model.generate_content(prompt)
        answer = response.text.strip() if hasattr(response, "text") else str(response)
    except Exception as e:
        answer = f"Error calling Gemini API: {e}"
    
    return answer

def extract_text_from_file(file_path):
    """Extract text from PDF, DOCX, or image files"""
    if not file_path:
        return None, None
    
    mime_type, _ = mimetypes.guess_type(file_path)
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.pdf' or (mime_type and 'pdf' in mime_type):
            text_content = []
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    num_pages = len(pdf_reader.pages)
                    print(f"PDF has {num_pages} pages")
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                            print(f"Extracted {len(page_text)} chars from page {page_num + 1}")
                    
                    extracted_text = "\n\n".join(text_content)
                    
                    if not extracted_text.strip():
                        print("PDF appears to be empty or scanned, treating as image")
                        return "pdf_image", file_path
                    
                    print(f"Total extracted text length: {len(extracted_text)}")
                    return "pdf", extracted_text
                    
            except Exception as pdf_error:
                print(f"PDF extraction error: {pdf_error}")
                return "pdf_image", file_path
        
        elif file_extension in ['.docx', '.doc'] or (mime_type and 'word' in mime_type):
            try:
                doc = Document(file_path)
                text_content = []
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_content.append(para.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            text_content.append(row_text)
                
                extracted_text = "\n\n".join(text_content)
                print(f"DOCX extracted text length: {len(extracted_text)}")
                
                if not extracted_text.strip():
                    return "docx", "Document appears to be empty"
                
                return "docx", extracted_text
                
            except Exception as docx_error:
                print(f"DOCX extraction error: {docx_error}")
                return "error", f"Could not read Word document: {docx_error}"
        
        elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'] or (mime_type and 'image' in mime_type):
            return "image", file_path
        
        else:
            return "unknown", None
    
    except Exception as e:
        print(f"Error extracting file content: {e}")
        return "error", str(e)

def chat_with_bot(message, file, history, selected_country, selected_category):
    """Handle chatbot interactions with RAG context and optional file upload"""
    if not message.strip() and file is None:
        return history, "", None
    
    context = ""
    if selected_country and selected_category:
        filtered_docs = [
            rule for rule in rules 
            if rule.get("country", "").lower() == selected_country.lower() 
            and rule.get("category", "").lower() == selected_category.lower()
        ]
        
        if not filtered_docs and index is not None and model is not None:
            try:
                search_query = {
                    "country": selected_country,
                    "category": selected_category.lower(),
                    "answers": {"question": message}
                }
                retrieved_docs = retrieve(search_query, rules, index, model)
                
                filtered_docs = [
                    doc for doc in retrieved_docs 
                    if doc.get("country", "").lower() == selected_country.lower() 
                    and doc.get("category", "").lower() == selected_category.lower()
                ]
                
                if not filtered_docs:
                    filtered_docs = [
                        doc for doc in retrieved_docs 
                        if doc.get("country", "").lower() == selected_country.lower()
                    ]
            except Exception as e:
                print(f"Chat RAG error: {e}")
        
        if filtered_docs:
            context_parts = []
            for doc in filtered_docs[:2]:
                context_parts.append(doc.get("text", ""))
            context = "\n\n".join(context_parts)
    
    if context and selected_country and selected_category:
        base_prompt = (
            f"You are a helpful visa assistant. Use this context about {selected_country} {selected_category} visa:\n\n"
            f"{context}\n\n"
        )
    else:
        base_prompt = (
            f"You are a helpful visa assistant. "
            f"Note: For more specific answers, please select a country and category above.\n\n"
        )
    
    try:
        if file is not None:
            file_type, content = extract_text_from_file(file)
            
            if file_type == "error":
                bot_message = f"Sorry, I couldn't read the file: {content}"
            elif file_type == "unknown":
                bot_message = "Sorry, this file type is not supported. Please upload PDF, DOCX, or image files."
            elif file_type == "image":
                pil_image = Image.open(content)
                if message.strip():
                    prompt = base_prompt + f"User question: {message}\n\nAnalyze the uploaded image and answer the question."
                    response = gemini_model.generate_content([prompt, pil_image])
                else:
                    prompt = base_prompt + "Analyze this document/image in the context of visa applications. What information does it contain and is it valid/sufficient?"
                    response = gemini_model.generate_content([prompt, pil_image])
                bot_message = response.text.strip() if hasattr(response, "text") else str(response)
            
            elif file_type == "pdf_image":
                
                bot_message = (
                    "This PDF appears to be a scanned document or image-based PDF. "
                    "The text extraction didn't work. Please try:\n"
                    "1. Converting it to a regular image (JPG/PNG) and uploading again\n"
                    "2. Or describing what information you need help with from this document"
                )
            
            elif file_type in ["pdf", "docx"]:
                
                if content and len(content.strip()) > 10:
                    
                    content_preview = content[:8000] if len(content) > 8000 else content
                    
                    if message.strip():
                        prompt = (
                            base_prompt +
                            f"User question: {message}\n\n"
                            f"Document content:\n{content_preview}\n\n"
                            f"Analyze this document and answer the question in the context of visa applications."
                        )
                    else:
                        prompt = (
                            base_prompt +
                            f"Document content:\n{content_preview}\n\n"
                            f"Analyze this document for visa application purposes. Extract key information and assess:\n"
                            f"1. What type of document is this?\n"
                            f"2. What key information does it contain?\n"
                            f"3. Is it relevant for visa applications?\n"
                            f"4. What's missing or what should be verified?"
                        )
                    response = gemini_model.generate_content(prompt)
                    bot_message = response.text.strip() if hasattr(response, "text") else str(response)
                else:
                    bot_message = "Sorry, I couldn't extract meaningful text from the document. It might be empty, corrupted, or image-based. Please try uploading it as an image format (JPG/PNG) instead."
        
        else:
            # No file, just text message
            prompt = base_prompt + f"User question: {message}\n\nProvide a helpful answer about visa applications."
            response = gemini_model.generate_content(prompt)
            bot_message = response.text.strip() if hasattr(response, "text") else str(response)
    
    except Exception as e:
        bot_message = f"Sorry, I encountered an error: {e}"
    
    
    user_message = message if message else "Uploaded a document"
    if file:
        file_name = os.path.basename(file)
        user_message += f" 📎 {file_name}"
    
    history.append({"role": "user", "content": user_message})
    history.append({"role": "assistant", "content": bot_message})
    return history, "", None

with gr.Blocks(
    theme=gr.themes.Soft(primary_hue="blue", secondary_hue="slate"),
    title="Visa Approval Probability Assistant",
    css="""
    #rag-sidebar {
        background-color: #f8f9fa;
        padding: 15px;
        border-radius: 8px;
        border: 1px solid #e0e0e0;
        height: 600px;
        overflow-y: auto;
    }
    #gemini-response-box {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 25px;
        border-radius: 12px;
        margin-top: 20px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    #gemini-response-box h3 {
        color: white;
        margin-top: 0;
        font-size: 1.3em;
        margin-bottom: 15px;
    }
    #gemini-response-content {
        background: rgba(255,255,255,0.95);
        color: #333;
        padding: 20px;
        border-radius: 8px;
        line-height: 1.6;
        max-height: 500px;
        overflow-y: auto;
    }
    .upload-section-title {
        margin-top: 20px;
        margin-bottom: 5px;
        font-size: 1.1em;
        color: #333;
    }
    .upload-hint {
        margin-top: 0;
        margin-bottom: 15px;
        font-size: 0.9em;
        color: #666;
        font-style: italic;
    }
    """
) as demo:
    gr.Markdown(
        """
        <div style="display:flex;align-items:center;gap:16px;margin-bottom:20px;">
            <img src="https://cdn-icons-png.flaticon.com/512/3062/3062634.png" width="48"/>
            <div>
                <h1 style="margin-bottom:0;">Visa Approval Probability Assistant</h1>
                <span style="font-size:1.1em;color:#555;">Gemini + RAG with Chatbot & Image Analysis</span>
            </div>
        </div>
        """
    )
    
    with gr.Row():
       
        with gr.Column(scale=3):
            gr.Markdown("## 🌍 Select Your Destination")
            country_dropdown = gr.Dropdown(
                choices=["Australia", "Canada", "USA"],
                label="Country",
                elem_id="country-dropdown"
            )
            
            category_dropdown = gr.Dropdown(
                choices=[],
                label="Visa Category",
                elem_id="category-dropdown",
                interactive=True
            )
            
            questions_column = gr.Column(visible=False)
            
            with questions_column:
                gr.Markdown("## 📝 Application Questions")
                question_components = []
                for i in range(10):
                    q = gr.Textbox(
                        label=f"Question {i+1}",
                        visible=False,
                        elem_id=f"question-{i}"
                    )
                    question_components.append(q)
                
                submit_btn = gr.Button("Submit Application", visible=False, variant="primary", size="lg")
            
           
            result_container = gr.Column(visible=False)
            with result_container:
                with gr.Group(elem_id="gemini-response-box"):
                    gr.Markdown("### Analysis")
                    result_box = gr.Markdown(elem_id="gemini-response-content")
            
            gr.Markdown("## 💬 Ask Questions", elem_id="chatbot-section")
            gr.Markdown("*Upload documents: **Images** (JPG, PNG), **PDFs**, or **Word documents** (.docx) for analysis*")
            chatbot = gr.Chatbot(
                label="Visa Assistant",
                height=350,
                show_label=True,
                type="messages"
            )
            with gr.Row():
                with gr.Column(scale=4):
                    chat_input = gr.Textbox(
                        placeholder="Ask me anything about visa requirements...",
                        show_label=False,
                        container=False
                    )
                with gr.Column(scale=1):
                    file_input = gr.File(
                        label="Upload Document",
                        file_types=["image", ".pdf", ".docx", ".doc"],
                        type="filepath"
                    )
            with gr.Row():
                chat_btn = gr.Button("Send", variant="primary", scale=1)
                clear_btn = gr.Button("Clear Chat", scale=1)
        
        with gr.Column(scale=2):
            rag_context_display = gr.Markdown(
                value="### 📋 RAG Context\n\nPlease select both **country** and **category** to view relevant visa requirements.",
                elem_id="rag-sidebar"
            )
    
    def update_categories(selected_country):
        """Update category dropdown when country is selected"""
        if not selected_country:
            return gr.Dropdown(choices=[], value=None), "### 📋 RAG Context\n\nPlease select a country first."
        cats = get_country_categories(selected_country)
        if not cats:
            return gr.Dropdown(choices=["No categories found"], value=None), f"### 📋 RAG Context\n\n**{selected_country}**\n\nNo categories found for this country."
        return gr.Dropdown(choices=cats, value=None), f"### 📋 RAG Context\n\n**{selected_country}**\n\nPlease select a category."
    
    def update_rag_sidebar(selected_country, selected_category):
        """Update RAG context sidebar"""
        return get_rag_context(selected_country, selected_category)
    
    def show_questions(selected_country, selected_category):
        """Show questions when category is selected"""
        if not selected_country or not selected_category:
            return [gr.Column(visible=False)] + [gr.Textbox(visible=False) for _ in range(10)] + [gr.Button(visible=False)]
        
        questions = build_questions(selected_country, selected_category.lower())
        updates = [gr.Column(visible=True)]
        
        for i in range(10):
            if i < len(questions):
                _, question_text = questions[i]
                updates.append(gr.Textbox(label=question_text, visible=True, value=""))
            else:
                updates.append(gr.Textbox(visible=False))
        
        updates.append(gr.Button(visible=True))
        return updates
    
    def submit_application(selected_country, selected_category, *question_answers):
        """Handle application submission"""
        result = process_application(selected_country, selected_category, *question_answers)
        return gr.Column(visible=True), result
    
    def clear_chat():
        """Clear chat history"""
        return [], "", None
    
    country_dropdown.change(
        update_categories,
        inputs=[country_dropdown],
        outputs=[category_dropdown, rag_context_display]
    )
    
    category_dropdown.change(
        show_questions,
        inputs=[country_dropdown, category_dropdown],
        outputs=[questions_column] + question_components + [submit_btn]
    )
    
    category_dropdown.change(
        update_rag_sidebar,
        inputs=[country_dropdown, category_dropdown],
        outputs=rag_context_display
    )
    
    submit_btn.click(
        submit_application,
        inputs=[country_dropdown, category_dropdown] + question_components,
        outputs=[result_container, result_box]
    )
    
    chat_btn.click(
        chat_with_bot,
        inputs=[chat_input, file_input, chatbot, country_dropdown, category_dropdown],
        outputs=[chatbot, chat_input, file_input]
    )
    
    chat_input.submit(
        chat_with_bot,
        inputs=[chat_input, file_input, chatbot, country_dropdown, category_dropdown],
        outputs=[chatbot, chat_input, file_input]
    )
    
    clear_btn.click(
        clear_chat,
        outputs=[chatbot, chat_input, file_input]
    )

if __name__ == "__main__":
    demo.launch()