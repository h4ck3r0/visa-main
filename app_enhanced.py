from flask import Flask, render_template, request, jsonify, session, send_file
import os
from dotenv import load_dotenv
import google.generativeai as genai
import sys
import json
from PIL import Image
import PyPDF2
from docx import Document
import mimetypes
import uuid
from werkzeug.utils import secure_filename
import base64
from io import BytesIO

from document_verifier import EnhancedDocumentVerifier

load_dotenv()

sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from rag_utils import prepare_rag_store, retrieve

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "your-secret-key-here")

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise RuntimeError("Set the GOOGLE_API_KEY environment variable.")

genai.configure(api_key=GOOGLE_API_KEY)
gemini_model = genai.GenerativeModel("gemini-2.5-flash")

print("Initializing Enhanced Document Verifier...")
enhanced_verifier = EnhancedDocumentVerifier()
print("Enhanced Document Verifier initialized")

json_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "visa_rules.json")
try:
    with open(json_path, 'r', encoding='utf-8') as f:
        rules = json.load(f)
    print(f"Loaded {len(rules)} rules from {json_path}")
except FileNotFoundError:
    print(f"Warning: {json_path} not found. Using empty rules.")
    rules = []


enable_rag = os.getenv('ENABLE_RAG', 'false').lower() == 'true'

embeddings, index, model = None, None, None
if enable_rag:
    try:
        print("Initializing RAG...")
        rules_rag, embeddings, index, model = prepare_rag_store()
        if len(rules_rag) > len(rules):
            rules = rules_rag
            print(f"Using {len(rules)} rules from RAG store")
    except Exception as e:
        print(f"Warning: Could not prepare RAG store: {e}")
        embeddings, index, model = None, None, None
else:
    print("RAG features disabled (set ENABLE_RAG=true to enable)")

UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_country_categories(country):
    if not country:
        return []
    cats = sorted(set(rule["category"] for rule in rules if rule.get("country") and rule["country"].lower() == country.lower()))
    cats = [c.title() for c in cats if c and c.strip()]
    return cats

def get_category_requirements(country, category):
    for rule in rules:
        rule_country = rule.get("country", "")
        rule_category = rule.get("category", "")
        if rule_country and rule_country.lower() == country.lower() and rule_category and rule_category.lower() == category.lower():
            if "requirements" in rule and isinstance(rule["requirements"], list):
                return rule["requirements"]
            req_text = rule.get("text", "")
            if "Required:" in req_text:
                reqs = req_text.split("Required:")[1].split(",")
            elif "Requirements:" in req_text:
                reqs = req_text.split("Requirements:")[1].split(",")
            else:
                reqs = []
            reqs = [r.strip().rstrip('.') for r in reqs if r.strip()]
            return reqs
    return []

def build_questions(country, category):
    reqs = get_category_requirements(country, category)
    
    if reqs:
        questions = []
        for i, req in enumerate(reqs):
            req_clean = req.strip()
            questions.append({"id": f"req_{i}", "question": f"Do you have: {req_clean}?"})
        return questions
    else:
        return [
            {"id": "passport_valid", "question": "Do you have a valid passport (valid for at least 6 months)?"},
            {"id": "financial_proof", "question": "Can you provide proof of sufficient funds (bank statements, payslips)?"},
            {"id": "travel_history", "question": "Do you have a good travel history (previous visas to other countries)?"},
            {"id": "employment_status", "question": "Are you currently employed/studying/retired with proof?"},
            {"id": "purpose_clear", "question": "Do you have a clear purpose of visit with supporting documents?"},
            {"id": "ties_home", "question": "Do you have strong ties to your home country (property, family, job)?"},
            {"id": "criminal_record", "question": "Do you have a clean criminal record with no visa rejections?"},
            {"id": "health_insurance", "question": "Do you have travel/health insurance coverage?"},
            {"id": "accommodation", "question": "Do you have confirmed accommodation/invitation letter?"},
            {"id": "return_ticket", "question": "Do you have a return ticket or travel itinerary?"}
        ]

def extract_text_from_file(file_path):
    """
    Legacy function maintained for compatibility
    New enhanced verification uses document_verifier.py
    """
    if not file_path:
        return None, None
    
    mime_type, _ = mimetypes.guess_type(file_path)
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.pdf' or (mime_type and 'pdf' in mime_type):
            text_content = []
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    
                    extracted_text = "\n\n".join(text_content)
                    if not extracted_text.strip():
                        return "pdf_image", file_path
                    return "pdf", extracted_text
            except Exception:
                return "pdf_image", file_path
        
        elif file_extension in ['.docx', '.doc'] or (mime_type and 'word' in mime_type):
            try:
                doc = Document(file_path)
                text_content = []
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_content.append(para.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            text_content.append(row_text)
                
                extracted_text = "\n\n".join(text_content)
                if not extracted_text.strip():
                    return "docx", "Document appears to be empty"
                return "docx", extracted_text
            except Exception as e:
                return "error", f"Could not read Word document: {e}"
        
        elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'] or (mime_type and 'image' in mime_type):
            return "image", file_path
        
        else:
            return "unknown", None
    
    except Exception as e:
        return "error", str(e)

def generate_enhanced_analysis_report(verification_results, message, country, category):
    """
    Generate comprehensive analysis report using enhanced verification results
    """
    try:
        score = verification_results.get('authenticity_score', 0)
        risk_level = verification_results.get('risk_level', 'UNKNOWN')
        
        if score >= 80:
            assessment = "✅ AUTHENTIC - Document appears genuine"
            recommendation_level = "APPROVED"
        elif score >= 60:
            assessment = "⚠️ QUESTIONABLE - Further verification recommended"
            recommendation_level = "REVIEW_REQUIRED"
        else:
            assessment = "❌ SUSPICIOUS - High probability of forgery"
            recommendation_level = "REJECTED"
        

        findings = []
        
        metadata = verification_results.get('metadata_analysis', {})
        if metadata.get('suspicious_indicators'):
            findings.append({
                'category': '📊 Metadata Analysis',
                'issues': metadata['suspicious_indicators'],
                'severity': 'medium'
            })
        
        content = verification_results.get('content_verification', {})
        
        expiry_validation = content.get('expiry_validation', {})
        if expiry_validation.get('suspicious_indicators'):
            findings.append({
                'category': '📅 Date & Timeline Analysis',
                'issues': expiry_validation['suspicious_indicators'],
                'severity': 'high'
            })
        
      
        validity_status = expiry_validation.get('validity_status', 'unknown')
        if validity_status == 'expired':
            findings.append({
                'category': '⏰ Document Validity',
                'issues': ['Document has expired'],
                'severity': 'high'
            })
        elif validity_status == 'suspicious_future_date':
            findings.append({
                'category': '⏰ Document Validity',
                'issues': ['Document contains suspicious future dates'],
                'severity': 'high'
            })
        
        
        spelling_grammar = content.get('spelling_grammar', {})
        spelling_errors = spelling_grammar.get('spelling_errors', [])
        if spelling_errors:
            error_words = [error['word'] for error in spelling_errors[:5]]  # Show first 5
            findings.append({
                'category': '📝 Spelling & Language Analysis',
                'issues': [f"Spelling errors found: {', '.join(error_words)}{'...' if len(spelling_errors) > 5 else ''}"],
                'severity': 'medium'
            })
        
        institutional_errors = spelling_grammar.get('institutional_errors', [])
        if institutional_errors:
            findings.append({
                'category': '🏛️ Institutional Terminology',
                'issues': [f"Incorrect institutional terms: {len(institutional_errors)} errors found"],
                'severity': 'high'
            })
        
        
        image_forensics = verification_results.get('image_forensics', {})
        if image_forensics.get('suspicious_indicators'):
            findings.append({
                'category': '🖼️ Image Forensics Analysis',
                'issues': image_forensics['suspicious_indicators'],
                'severity': 'high'
            })
        
        
        ml_analysis = verification_results.get('ml_analysis', {})
        if ml_analysis.get('anomaly_score', 0) > 0.7:
            findings.append({
                'category': '🤖 Machine Learning Analysis',
                'issues': ['Document patterns indicate high probability of forgery'],
                'severity': 'high'
            })
        
        
        report_html = f"""
        <div class="verification-report">
            <div class="report-header">
                <h2>📋 Enhanced Document Verification Report</h2>
                <div class="assessment-badge {recommendation_level.lower()}">{assessment}</div>
            </div>
            
            <div class="score-section">
                <div class="score-circle">
                    <span class="score-number">{score:.1f}</span>
                    <span class="score-label">Authenticity Score</span>
                </div>
                <div class="risk-info">
                    <strong>Risk Level:</strong> {risk_level}<br>
                    <strong>Recommendation:</strong> {recommendation_level.replace('_', ' ').title()}
                </div>
            </div>
        """
        
        if findings:
            report_html += "<div class='findings-section'><h3>🔍 Detailed Findings:</h3>"
            for finding in findings:
                severity_class = finding['severity']
                issues_html = '<br>'.join([f"• {issue}" for issue in finding['issues']])
                report_html += f"""
                <div class='finding-item {severity_class}'>
                    <strong>{finding['category']}</strong><br>
                    {issues_html}
                </div>
                """
            report_html += "</div>"
        else:
            report_html += "<div class='no-issues'>✅ No significant issues detected</div>"
        
        
        recommendations = verification_results.get('recommendations', [])
        if recommendations:
            report_html += "<div class='recommendations-section'><h3>💡 Recommendations:</h3><ul>"
            for rec in recommendations:
                report_html += f"<li>{rec}</li>"
            report_html += "</ul></div>"
        
        
        report_html += f"""
        <details class="technical-details">
            <summary>🔧 Technical Analysis Details</summary>
            <div class="tech-content">
                <h4>File Information:</h4>
                <ul>
                    <li><strong>Filename:</strong> {verification_results.get('file_info', {}).get('filename', 'Unknown')}</li>
                    <li><strong>File Size:</strong> {verification_results.get('file_info', {}).get('size_bytes', 0):,} bytes</li>
                    <li><strong>MIME Type:</strong> {verification_results.get('file_info', {}).get('mime_type', 'Unknown')}</li>
                </ul>
                
                <h4>Analysis Components:</h4>
                <ul>
                    <li>✓ Metadata Analysis</li>
                    <li>✓ Content Verification</li>
                    <li>✓ Date & Timeline Validation</li>
                    <li>✓ Spelling & Grammar Check</li>
                    <li>✓ Image Forensics (if applicable)</li>
                    <li>✓ ML-based Anomaly Detection</li>
                </ul>
            </div>
        </details>
        """
        
        report_html += "</div>"
        
        return report_html
        
    except Exception as e:
        return f"""
        <div class="error-report">
            <h3>⚠️ Analysis Error</h3>
            <p>An error occurred while generating the detailed report: {str(e)}</p>
            <p>Basic score: {verification_results.get('authenticity_score', 0):.1f}/100</p>
        </div>
        """

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/application')
def application():
    return render_template('application.html')

@app.route('/chat')
def chat():
    return render_template('chat.html')

@app.route('/documents')
def documents():
    return render_template('documents.html')

@app.route('/help')
def help_page():
    return render_template('help.html')

@app.route('/api/countries')
def get_countries():
    countries = sorted(set(rule["country"] for rule in rules if rule.get("country")))
    return jsonify(countries)

@app.route('/api/categories/<country>')
def get_categories(country):
    categories = get_country_categories(country)
    return jsonify(categories)

@app.route('/api/questions/<country>/<category>')
def get_questions(country, category):
    questions = build_questions(country, category)
    return jsonify(questions)

@app.route('/api/rag-context/<country>/<category>')
def get_rag_context(country, category):
    if not country or not category:
        return jsonify({"context": "Please select both country and category to view relevant visa requirements."})
    
    filtered_docs = [
        rule for rule in rules 
        if rule.get("country", "").lower() == country.lower() 
        and rule.get("category", "").lower() == category.lower()
    ]
    
    if not filtered_docs and index is not None and model is not None:
        try:
            search_query = {
                "country": country,
                "category": category.lower(),
                "answers": {}
            }
            retrieved_docs = retrieve(search_query, rules, index, model)
            
            filtered_docs = [
                doc for doc in retrieved_docs 
                if doc.get("country", "").lower() == country.lower() 
                and doc.get("category", "").lower() == category.lower()
            ]
            
            if not filtered_docs:
                filtered_docs = [
                    doc for doc in retrieved_docs 
                    if doc.get("country", "").lower() == country.lower()
                ]
        except Exception as e:
            print(f"RAG retrieval error: {e}")
    
    if not filtered_docs:
        return jsonify({"context": f"No specific requirements found for {country} - {category.title()}."})
    
    context_html = f"<h3>{country} - {category.title()}</h3>"
    for doc in filtered_docs:
        doc_title = doc.get("title", "Unknown")
        doc_text = doc.get("text", "No details available")
        
        context_html += f"<div class='requirement-card'><h4>{doc_title}</h4>"
        
        if "requirements" in doc and isinstance(doc["requirements"], list):
            context_html += "<ul>"
            for req in doc["requirements"]:
                context_html += f"<li>{req}</li>"
            context_html += "</ul>"
        
        context_html += f"<p>{doc_text}</p></div>"
    
    return jsonify({"context": context_html})

@app.route('/api/submit-application', methods=['POST'])
def submit_application():
    data = request.json
    country = data.get('country')
    category = data.get('category')
    answers = data.get('answers', {})
    
    # Process application similar to original logic
    search_query = {
        "country": country,
        "category": category.lower(),
        "answers": answers
    }
    
    retrieved_docs = retrieve(search_query, rules, index, model) if (index and model) else []
    
    filtered_docs = [
        doc for doc in retrieved_docs 
        if doc.get("country", "").lower() == country.lower() 
        and doc.get("category", "").lower() == category.lower()
    ]
    
    if not filtered_docs:
        filtered_docs = [
            doc for doc in retrieved_docs 
            if doc.get("country", "").lower() == country.lower()
        ]
    
    # Format prompt
    context = ""
    if filtered_docs:
        context_parts = []
        for doc in filtered_docs:
            doc_country = doc.get("country", "Unknown")
            doc_category = doc.get("category", "Unknown")
            doc_text = doc.get("text", "")
            context_parts.append(f"[{doc_country} - {doc_category}]\n{doc_text}")
        context = "\n\n".join(context_parts)
    
    answer_text = "\n".join([f"- {k}: {v}" for k, v in answers.items()])
    
    if context:
        prompt = (
            f"You are a visa application expert. Use the following official visa requirements as context:\n\n"
            f"{context}\n\n"
            f"IMPORTANT: The applicant is applying for a {category.upper()} visa to {country.upper()}. "
            f"Only use requirements and information specific to {country} {category} visa.\n\n"
            f"Visa Application Information:\n"
            f"{answer_text}\n\n"
            f"Based on this information:\n"
            f"1. Estimate the visa approval probability (as a percentage)\n"
            f"2. List any missing documents or risk factors\n"
            f"3. Provide professional advice to improve approval chances\n"
            f"4. Be specific to {country} {category} visa requirements"
        )
    else:
        prompt = (
            f"You are a visa application expert. Based on your knowledge:\n\n"
            f"The applicant is applying for a {category.upper()} visa to {country.upper()}.\n\n"
            f"Visa Application Information:\n"
            f"{answer_text}\n\n"
            f"Based on this information:\n"
            f"1. Estimate the visa approval probability (as a percentage)\n"
            f"2. List any missing documents or risk factors\n"
            f"3. Provide professional advice to improve approval chances\n"
            f"4. Be specific to {country} {category} visa requirements"
        )
    
    try:
        response = gemini_model.generate_content(prompt)
        result = response.text.strip() if hasattr(response, "text") else str(response)
        return jsonify({"success": True, "result": result})
    except Exception as e:
        return jsonify({"success": False, "error": f"Error calling Gemini API: {e}"})

@app.route('/api/chat', methods=['POST'])
def chat_api():
    data = request.json
    message = data.get('message', '')
    country = data.get('country')
    category = data.get('category')
    
    # Get context if country and category are provided
    context = ""
    if country and category:
        filtered_docs = [
            rule for rule in rules 
            if rule.get("country", "").lower() == country.lower() 
            and rule.get("category", "").lower() == category.lower()
        ]
        
        if not filtered_docs and index is not None and model is not None:
            try:
                search_query = {
                    "country": country,
                    "category": category.lower(),
                    "answers": {"question": message}
                }
                retrieved_docs = retrieve(search_query, rules, index, model)
                
                filtered_docs = [
                    doc for doc in retrieved_docs 
                    if doc.get("country", "").lower() == country.lower() 
                    and doc.get("category", "").lower() == category.lower()
                ]
                
                if not filtered_docs:
                    filtered_docs = [
                        doc for doc in retrieved_docs 
                        if doc.get("country", "").lower() == country.lower()
                    ]
            except Exception as e:
                print(f"Chat RAG error: {e}")
        
        if filtered_docs:
            context_parts = []
            for doc in filtered_docs[:2]:
                context_parts.append(doc.get("text", ""))
            context = "\n\n".join(context_parts)
    
    if context and country and category:
        base_prompt = (
            f"You are a helpful visa assistant. Use this context about {country} {category} visa:\n\n"
            f"{context}\n\n"
        )
    else:
        base_prompt = (
            f"You are a helpful visa assistant. "
            f"Note: For more specific answers, please select a country and category.\n\n"
        )
    
    try:
        prompt = base_prompt + f"User question: {message}\n\nProvide a helpful answer about visa applications."
        response = gemini_model.generate_content(prompt)
        bot_message = response.text.strip() if hasattr(response, "text") else str(response)
        return jsonify({"success": True, "message": bot_message})
    except Exception as e:
        return jsonify({"success": False, "error": f"Error: {e}"})

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """
    Enhanced file upload with comprehensive document verification
    """
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "No file provided"})
    
    file = request.files['file']
    message = request.form.get('message', '')
    country = request.form.get('country')
    category = request.form.get('category')
    
    if file.filename == '':
        return jsonify({"success": False, "error": "No file selected"})
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        unique_filename = f"{uuid.uuid4()}_{filename}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(file_path)
        
        try:
            # Use enhanced document verification system
            print(f"Starting enhanced verification for: {filename}")
            verification_results = enhanced_verifier.verify_document(file_path)
            
            # Clean up uploaded file
            os.remove(file_path)
            
            # Check if verification was successful
            if 'error' in verification_results:
                return jsonify({
                    "success": False, 
                    "error": f"Verification failed: {verification_results['error']}"
                })
            
            # Generate enhanced analysis report
            analysis_report = generate_enhanced_analysis_report(
                verification_results, message, country, category
            )
            
            # Generate AI-powered insights using Gemini with detailed verification results
            ai_final_assessment = ""
            ai_forgery_detection = ""
            
            try:
                # Build comprehensive context for Gemini
                metadata_analysis = verification_results.get('metadata_analysis', {})
                content_analysis = verification_results.get('content_verification', {})
                image_analysis = verification_results.get('image_forensics', {})
                
                # Prepare detailed findings
                suspicious_indicators = []
                suspicious_indicators.extend(metadata_analysis.get('suspicious_indicators', []))
                suspicious_indicators.extend(content_analysis.get('spelling_grammar', {}).get('suspicious_indicators', []))
                
                expiry_info = content_analysis.get('expiry_validation', {})
                if expiry_info.get('suspicious_indicators'):
                    suspicious_indicators.extend(expiry_info['suspicious_indicators'])
                
                if image_analysis.get('suspicious_indicators'):
                    suspicious_indicators.extend(image_analysis['suspicious_indicators'])
                
                # Create concise Gemini prompt for FINAL ASSESSMENT
                final_assessment_prompt = f"""INSTRUCTIONS: Reply with EXACTLY 2 sentences. NO special characters, NO **, NO formatting.

Score is {verification_results['authenticity_score']}/100. Risk is {verification_results['risk_level']}.

Is document authentic or fake? Answer ONLY based on score: above 80=authentic, below 60=fake. Keep plain text, max 25 words."""
                
                # Get final assessment from Gemini
                final_response = gemini_model.generate_content(final_assessment_prompt)
                ai_final_assessment = final_response.text.strip() if hasattr(final_response, "text") else ""
                
                # Create concise prompt for FORGERY DETECTION ANALYSIS
                forgery_detection_prompt = f"""INSTRUCTIONS: Reply with EXACTLY 2 sentences. NO special characters, NO **, NO formatting.

Score is {verification_results['authenticity_score']}/100. Risk is {verification_results['risk_level']}.

Based ONLY on score: above 80=NO forgery risk, below 60=high risk. Is forgery likely or not? Keep plain text, max 25 words."""
                
                # Get forgery detection analysis from Gemini
                forgery_response = gemini_model.generate_content(forgery_detection_prompt)
                ai_forgery_detection = forgery_response.text.strip() if hasattr(forgery_response, "text") else ""
                
                # Validate responses - if contradictory, fix them based on actual score
                score = verification_results['authenticity_score']
                if score >= 80:
                    if 'high' in ai_forgery_detection.lower() and 'risk' in ai_forgery_detection.lower():
                        ai_forgery_detection = "Forgery risk is low. High authenticity score indicates document is genuine."
                elif score <= 60:
                    if 'no' in ai_final_assessment.lower() or 'not' in ai_final_assessment.lower():
                        ai_final_assessment = "Document appears questionable. Low score suggests potential issues."
                
            except Exception as e:
                ai_final_assessment = f"AI analysis unavailable: {str(e)}"
                ai_forgery_detection = f"Forgery detection analysis unavailable: {str(e)}"
            
            return jsonify({
                "success": True,
                "enhanced_verification": True,
                "authenticity_score": verification_results['authenticity_score'],
                "risk_level": verification_results['risk_level'],
                "analysis_report": analysis_report,
                "ai_final_assessment": ai_final_assessment,
                "ai_forgery_detection": ai_forgery_detection,
                "detailed_results": verification_results,
                "recommendations": verification_results.get('recommendations', []),
                "file_type": "enhanced_analysis"
            })
            
        except Exception as e:
            # Clean up file if it exists
            if os.path.exists(file_path):
                os.remove(file_path)
            
            print(f"Enhanced verification error: {str(e)}")
            
            # Fallback to legacy verification
            try:
                file.seek(0)  # Reset file pointer
                file.save(file_path)
                
                file_type, content = extract_text_from_file(file_path)
                os.remove(file_path)
                
                if file_type == "error":
                    return jsonify({"success": False, "error": f"Could not read file: {content}"})
                elif file_type == "unknown":
                    return jsonify({"success": False, "error": "Unsupported file type"})
                elif file_type == "image":
                    # Handle image analysis with Gemini
                    pil_image = Image.open(BytesIO(file.read()))
                    
                    base_prompt = "You are a document verification assistant. "
                    if country and category:
                        base_prompt += f"Focus on {country} {category} visa requirements. "
                    
                    if message.strip():
                        prompt = base_prompt + f"User question: {message}\n\nAnalyze the uploaded image for document authenticity and answer the question."
                    else:
                        prompt = base_prompt + "Analyze this document/image for authenticity. Look for signs of forgery, tampering, or inconsistencies."
                    
                    response = gemini_model.generate_content([prompt, pil_image])
                    bot_message = response.text.strip() if hasattr(response, "text") else str(response)
                    
                    return jsonify({
                        "success": True, 
                        "message": bot_message, 
                        "file_type": "image",
                        "enhanced_verification": False,
                        "fallback_mode": True
                    })
                
                elif file_type == "pdf_image":
                    return jsonify({
                        "success": True, 
                        "message": "This PDF appears to be a scanned document. The enhanced verification system detected it but could not complete full analysis. Please try converting it to an image format (JPG/PNG) for better analysis.",
                        "file_type": "pdf_image",
                        "enhanced_verification": False,
                        "fallback_mode": True
                    })
                
                elif file_type in ["pdf", "docx"]:
                    if content and len(content.strip()) > 10:
                        content_preview = content[:8000] if len(content) > 8000 else content
                        
                        base_prompt = "You are a document verification assistant. "
                        if country and category:
                            base_prompt += f"Focus on {country} {category} visa requirements. "
                        
                        if message.strip():
                            prompt = (
                                base_prompt +
                                f"User question: {message}\n\n"
                                f"Document content:\n{content_preview}\n\n"
                                f"Analyze this document for authenticity and answer the question. Look for inconsistencies, spelling errors, or suspicious content."
                            )
                        else:
                            prompt = (
                                base_prompt +
                                f"Document content:\n{content_preview}\n\n"
                                f"Analyze this document for visa application purposes and authenticity. Assess:\n"
                                f"1. Document type and relevance\n"
                                f"2. Content quality and consistency\n"
                                f"3. Potential signs of forgery or tampering\n"
                                f"4. Recommendations for visa application"
                            )
                        
                        response = gemini_model.generate_content(prompt)
                        bot_message = response.text.strip() if hasattr(response, "text") else str(response)
                        
                        return jsonify({
                            "success": True, 
                            "message": bot_message, 
                            "file_type": file_type,
                            "enhanced_verification": False,
                            "fallback_mode": True
                        })
                    else:
                        return jsonify({
                            "success": False, 
                            "error": "Could not extract meaningful text from the document."
                        })
                        
            except Exception as fallback_error:
                return jsonify({
                    "success": False, 
                    "error": f"Both enhanced and fallback verification failed: {str(fallback_error)}"
                })
    
    return jsonify({"success": False, "error": "Invalid file type"})

@app.route('/api/verification-status')
def verification_status():
    """Get the status of the enhanced verification system"""
    try:
        status = {
            "enhanced_verification_available": True,
            "capabilities": [
                "Metadata Analysis",
                "Date & Timeline Validation", 
                "Spelling & Grammar Check",
                "Image Forensics Analysis",
                "ML-based Anomaly Detection",
                "Comprehensive Authenticity Scoring"
            ],
            "supported_formats": ["PDF", "DOCX", "JPG", "PNG", "TIFF"],
            "version": "2.0"
        }
        return jsonify(status)
    except:
        return jsonify({
            "enhanced_verification_available": False,
            "fallback_mode": True,
            "version": "1.0"
        })

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)